import os
import json
import requests
import time
import tempfile
import re
import statistics

from flask import Flask, request

from PIL import (
    Image,
    ImageEnhance,
    ImageFilter,
    ImageOps
)

import pytesseract

from pytesseract import Output

from io import BytesIO

from docx import Document

from docx.shared import Pt, Cm

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.table import (
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT
)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


app = Flask(__name__)


# =========================================================
# НАСТРОЙКИ
# =========================================================

BOT_TOKEN = os.environ.get("BOT_TOKEN")

API = f"https://api.telegram.org/bot{BOT_TOKEN}"


# =========================================================
# ПАМЯТЬ БОТА
# =========================================================

processed_updates = set()

user_texts = {}

user_ocr_data = {}

photo_processing = {}


# =========================================================
# ЗАЩИТА ОТ ПОВТОРНЫХ UPDATE
# =========================================================

def already_processed(update_id):

    if update_id in processed_updates:

        print(
            "⚠️ UPDATE УЖЕ ОБРАБОТАН:",
            update_id,
            flush=True
        )

        return True

    processed_updates.add(update_id)

    if len(processed_updates) > 1000:

        oldest = next(
            iter(processed_updates)
        )

        processed_updates.discard(
            oldest
        )

    return False


# =========================================================
# TELEGRAM API
# =========================================================

def telegram(method, data=None):

    try:

        response = requests.post(
            f"{API}/{method}",
            data=data or {},
            timeout=30
        )

        print(
            "Telegram:",
            method,
            response.status_code,
            response.text,
            flush=True
        )

        return response.json()

    except Exception as error:

        print(
            "TELEGRAM ERROR:",
            error,
            flush=True
        )

        return {
            "ok": False
        }


# =========================================================
# ОТПРАВКА СООБЩЕНИЯ
# =========================================================

def send_message(
    chat_id,
    text,
    keyboard=None
):

    data = {
        "chat_id": chat_id,
        "text": text
    }

    if keyboard is not None:

        data["reply_markup"] = json.dumps(
            keyboard,
            ensure_ascii=False
        )

    return telegram(
        "sendMessage",
        data
    )


# =========================================================
# ДЛИННЫЙ ТЕКСТ
# =========================================================

def send_long_message(
    chat_id,
    text
):

    max_length = 4000

    if len(text) <= max_length:

        send_message(
            chat_id,
            text
        )

        return

    for i in range(
        0,
        len(text),
        max_length
    ):

        send_message(
            chat_id,
            text[
                i:i + max_length
            ]
        )


# =========================================================
# СКАЧИВАНИЕ ФОТО
# =========================================================

def download_telegram_photo(
    file_id
):

    print(
        "Получаем файл:",
        file_id,
        flush=True
    )

    result = telegram(
        "getFile",
        {
            "file_id": file_id
        }
    )

    if not result.get("ok"):

        print(
            "GET FILE ERROR:",
            result,
            flush=True
        )

        return None

    file_path = result[
        "result"
    ][
        "file_path"
    ]

    url = (
        f"https://api.telegram.org/file/"
        f"bot{BOT_TOKEN}/{file_path}"
    )

    try:

        response = requests.get(
            url,
            timeout=30
        )

        if response.status_code != 200:

            print(
                "DOWNLOAD ERROR:",
                response.status_code,
                flush=True
            )

            return None

        return response.content

    except Exception as error:

        print(
            "DOWNLOAD EXCEPTION:",
            error,
            flush=True
        )

        return None


# =========================================================
# ПОДГОТОВКА ИЗОБРАЖЕНИЯ
# =========================================================

def prepare_image(image):

    print(
        "🖼️ ПОДГОТОВКА ИЗОБРАЖЕНИЯ",
        flush=True
    )

    image = image.convert(
        "RGB"
    )

    width, height = image.size

    print(
        "Исходный размер:",
        width,
        "x",
        height,
        flush=True
    )

    # -----------------------------------------------------
    # Увеличиваем маленькие изображения
    # -----------------------------------------------------

    target_width = 1800

    if width < target_width:

        scale = target_width / width

        image = image.resize(
            (
                int(width * scale),
                int(height * scale)
            ),
            Image.Resampling.LANCZOS
        )

    # -----------------------------------------------------
    # Ограничиваем максимальный размер
    # -----------------------------------------------------

    max_width = 2400

    if image.width > max_width:

        scale = max_width / image.width

        image = image.resize(
            (
                max_width,
                int(image.height * scale)
            ),
            Image.Resampling.LANCZOS
        )

    # -----------------------------------------------------
    # Серый
    # -----------------------------------------------------

    image = ImageOps.grayscale(
        image
    )

    # -----------------------------------------------------
    # Автоконтраст
    # -----------------------------------------------------

    image = ImageOps.autocontrast(
        image
    )

    # -----------------------------------------------------
    # Контраст
    # -----------------------------------------------------

    image = ImageEnhance.Contrast(
        image
    ).enhance(1.35)

    # -----------------------------------------------------
    # Резкость
    # -----------------------------------------------------

    image = image.filter(
        ImageFilter.SHARPEN
    )

    print(
        "Размер после подготовки:",
        image.size,
        flush=True
    )

    return image


# =========================================================
# ОЧИСТКА ОДНОГО OCR-СЛОВА
# =========================================================

def clean_ocr_word(word):

    if not word:
        return ""

    word = word.replace(
        "\x0c",
        ""
    )

    # Убираем только явно мусорные управляющие символы.
    # Пунктуацию НЕ трогаем.

    word = "".join(
        char
        for char in word
        if char.isprintable()
    )

    return word.strip()


# =========================================================
# OCR С КООРДИНАТАМИ
# =========================================================

def recognize_document(
    image
):

    print(
        "🧠 DOCUMENT OCR START",
        flush=True
    )

    # -----------------------------------------------------
    # Основной режим
    # -----------------------------------------------------

    data = pytesseract.image_to_data(
        image,
        lang="rus+eng",
        config="--oem 3 --psm 3",
        output_type=Output.DICT
    )

    words = []

    total = len(
        data.get("text", [])
    )

    for i in range(total):

        raw_text = data["text"][i]

        text = clean_ocr_word(
            raw_text
        )

        if not text:
            continue

        try:

            confidence = float(
                data["conf"][i]
            )

        except Exception:

            confidence = -1

        if confidence < 15:
            continue

        try:

            left = int(
                data["left"][i]
            )

            top = int(
                data["top"][i]
            )

            width = int(
                data["width"][i]
            )

            height = int(
                data["height"][i]
            )

            block = int(
                data["block_num"][i]
            )

            paragraph = int(
                data["par_num"][i]
            )

            line = int(
                data["line_num"][i]
            )

        except Exception:

            continue

        words.append({

            "text": text,

            "left": left,

            "top": top,

            "right": left + width,

            "bottom": top + height,

            "width": width,

            "height": height,

            "confidence": confidence,

            "block": block,

            "paragraph": paragraph,

            "line": line

        })

    print(
        "OCR WORDS:",
        len(words),
        flush=True
    )

    if not words:
        return None

    # -----------------------------------------------------
    # Группируем слова в визуальные строки
    # -----------------------------------------------------

    lines = group_words_into_lines(
        words
    )

    print(
        "OCR LINES:",
        len(lines),
        flush=True
    )

    return {
        "words": words,
        "lines": lines
    }


# =========================================================
# ГРУППИРОВКА СЛОВ В СТРОКИ
# =========================================================

def group_words_into_lines(
    words
):

    if not words:
        return []

    words = sorted(
        words,
        key=lambda item: (
            item["top"],
            item["left"]
        )
    )

    lines = []

    current = []

    current_top = None

    heights = [
        word["height"]
        for word in words
        if word["height"] > 0
    ]

    if heights:

        median_height = statistics.median(
            heights
        )

    else:

        median_height = 20

    tolerance = max(
        8,
        int(median_height * 0.65)
    )

    for word in words:

        if not current:

            current = [word]

            current_top = word["top"]

            continue

        if abs(
            word["top"] - current_top
        ) <= tolerance:

            current.append(
                word
            )

        else:

            current = sorted(
                current,
                key=lambda item: item["left"]
            )

            lines.append(
                current
            )

            current = [word]

            current_top = word["top"]

    if current:

        current = sorted(
            current,
            key=lambda item: item["left"]
        )

        lines.append(
            current
        )

    return lines


# =========================================================
# СОЕДИНЕНИЕ СЛОВ В СТРОКУ
# =========================================================

def line_to_text(
    line
):

    if not line:
        return ""

    result = ""

    previous = None

    for word in line:

        text = word["text"]

        if previous is None:

            result = text

        else:

            gap = (
                word["left"]
                -
                previous["right"]
            )

            # Если расстояние большое —
            # обычный пробел.
            #
            # Мы не используем несколько пробелов
            # как основу структуры. Структура берётся
            # из координат.

            if gap > 2:

                result += " "

            result += text

        previous = word

    return result.strip()


# =========================================================
# СОЗДАНИЕ ТЕКСТА ИЗ OCR
# =========================================================

def ocr_to_text(
    ocr_data
):

    if not ocr_data:
        return ""

    lines = ocr_data.get(
        "lines",
        []
    )

    result = []

    for line in lines:

        text = line_to_text(
            line
        )

        if text:

            result.append(
                text
            )

    return "\n".join(
        result
    ).strip()


# =========================================================
# ОЧИСТКА OCR ТЕКСТА
# =========================================================

def clean_ocr_text(
    text
):

    if not text:
        return ""

    result = []

    for raw_line in text.splitlines():

        # Не делаем strip агрессивно.
        # Убираем только лишние крайние пробелы.

        line = raw_line.strip()

        if not line:

            result.append("")

            continue

        # Повторяющиеся пробелы внутри текста
        # можно сократить.
        #
        # Но пунктуацию не изменяем.

        line = re.sub(
            r"[ \t]{2,}",
            " ",
            line
        )

        result.append(
            line
        )

    # Не больше двух пустых строк подряд.

    final = []

    empty_count = 0

    for line in result:

        if not line:

            empty_count += 1

            if empty_count <= 2:

                final.append("")

        else:

            empty_count = 0

            final.append(
                line
            )

    return "\n".join(
        final
    ).strip()


# =========================================================
# ОПРЕДЕЛЕНИЕ ЗАГОЛОВКА
# =========================================================

def is_heading(
    line
):

    line = line.strip()

    if not line:
        return False

    words = line.split()

    if len(words) > 10:
        return False

    if len(line) > 90:
        return False

    # Никогда не считаем пункт списка заголовком

    if re.match(
        r"^(\d+[\.\)]|\d+\.\d+[\.\)]|[-•*–])\s+",
        line
    ):

        return False

    # Если строка заканчивается
    # обычной пунктуацией — скорее всего текст.

    if line.endswith(
        (".", ",", ";", ":", "?", "!")
    ):

        return False

    letters = [
        char
        for char in line
        if char.isalpha()
    ]

    if not letters:
        return False

    uppercase = [
        char
        for char in letters
        if char.isupper()
    ]

    ratio = (
        len(uppercase)
        /
        len(letters)
    )

    # Явный CAPS-заголовок

    if (
        ratio >= 0.80
        and len(words) <= 8
    ):

        return True

    # Очень короткая строка,
    # но только если она визуально похожа
    # на отдельный заголовок.

    if (
        len(words) <= 5
        and len(line) <= 55
        and line[0].isupper()
    ):

        # Не превращаем обычные предложения
        # в заголовки.

        if not re.search(
            r"[,.!?;:]",
            line
        ):

            return True

    return False


# =========================================================
# СПИСКИ
# =========================================================

def is_list_item(
    line
):

    return bool(
        re.match(
            r"^(\d+[\.\)]|\d+\.\d+[\.\)]|[-•*–])\s+",
            line
        )
    )


def is_numbered_item(line):

    return (
        bool(
            re.match(
                r"^\d+[\.\)]\s+",
                line
            )
        )
        or
        bool(
            re.match(
                r"^\d+\.\d+[\.\)]\s+",
                line
            )
        )
    )
# =========================================================
# АНАЛИЗ КОЛОНОК
# =========================================================

def get_line_gaps(
    line
):

    if len(line) < 2:
        return []

    gaps = []

    for i in range(
        1,
        len(line)
    ):

        previous = line[i - 1]

        current = line[i]

        gap = (
            current["left"]
            -
            previous["right"]
        )

        if gap > 0:

            gaps.append({
                "position": current["left"],
                "gap": gap
            })

    return gaps


# =========================================================
# ОПРЕДЕЛЕНИЕ ТАБЛИЦЫ ПО КООРДИНАТАМ
# =========================================================

def detect_table_blocks(lines):

    if not lines or len(lines) < 3:
        return []

    blocks = []

    # Ищем последовательности строк,
    # в которых визуально присутствуют 2+ колонки.
    candidate_indexes = []

    for index, line in enumerate(lines):

        if len(line) < 2:
            continue

        gaps = get_line_gaps(line)

        large_gaps = [
            gap
            for gap in gaps
            if gap["gap"] >= 70
        ]

        if len(large_gaps) >= 1:
            candidate_indexes.append(index)

    if not candidate_indexes:
        return []

    current = [candidate_indexes[0]]

    for index in candidate_indexes[1:]:

        if index - current[-1] <= 1:

            current.append(index)

        else:

            if len(current) >= 2:
                blocks.append(current)

            current = [index]

    if len(current) >= 2:
        blocks.append(current)

    return blocks

# =========================================================
# ПОЛУЧЕНИЕ КОЛОНОК ТАБЛИЦЫ
# =========================================================

def build_table_from_lines(lines):

    if not lines or len(lines) < 2:
        return None

    # -----------------------------------------------------
    # Определяем реальные вертикальные позиции слов
    # -----------------------------------------------------

    all_words = []

    for line in lines:

        for word in line:

            all_words.append(word)

    if len(all_words) < 4:
        return None

    # -----------------------------------------------------
    # Ищем большие горизонтальные разрывы.
    # Они обычно разделяют колонки.
    # -----------------------------------------------------

    gaps = []

    for line in lines:

        sorted_line = sorted(
            line,
            key=lambda item: item["left"]
        )

        for i in range(1, len(sorted_line)):

            previous = sorted_line[i - 1]
            current = sorted_line[i]

            gap = (
                current["left"]
                -
                previous["right"]
            )

            if gap >= 70:

                gaps.append({
                    "x": current["left"],
                    "gap": gap
                })

    if not gaps:
        return None

    # -----------------------------------------------------
    # Группируем похожие позиции начала колонок
    # -----------------------------------------------------

    positions = sorted(
        gap["x"]
        for gap in gaps
    )

    clusters = []

    for position in positions:

        if not clusters:

            clusters.append([position])

            continue

        average = statistics.mean(
            clusters[-1]
        )

        if abs(position - average) <= 100:

            clusters[-1].append(position)

        else:

            clusters.append([position])

    column_starts = []

    for cluster in clusters:

        if len(cluster) >= 1:

            column_starts.append(
                int(
                    statistics.median(cluster)
                )
            )

    # -----------------------------------------------------
    # Для нормальной таблицы достаточно 2 колонок
    # -----------------------------------------------------

    if len(column_starts) < 2:
        return None

    # Не допускаем слишком много ложных колонок

    column_starts = column_starts[:4]

    rows = []

    # -----------------------------------------------------
    # Раскладываем слова по колонкам
    # -----------------------------------------------------

    for line in lines:

        cells = [
            []
            for _ in column_starts
        ]

        for word in sorted(
            line,
            key=lambda item: item["left"]
        ):

            x = word["left"]

            # Последняя колонка, начало которой
            # находится левее слова.
            column_index = 0

            for i, start in enumerate(
                column_starts
            ):

                if x >= start:
                    column_index = i

            cells[column_index].append(
                word["text"]
            )

        row = [
            " ".join(cell).strip()
            for cell in cells
        ]

        # Убираем полностью пустые строки

        if any(row):

            rows.append(row)

    if len(rows) < 2:
        return None

    # -----------------------------------------------------
    # Проверяем, действительно ли это таблица
    # -----------------------------------------------------

    multi_column_rows = 0

    for row in rows:

        filled = sum(
            1
            for cell in row
            if cell.strip()
        )

        if filled >= 2:
            multi_column_rows += 1

    # Если меньше двух строк имеют несколько колонок,
    # скорее всего это обычный текст.

    if multi_column_rows < 2:
        return None

    return rows


# =========================================================
# ДОБАВЛЕНИЕ ТАБЛИЦЫ
# =========================================================

def add_table(
    document,
    rows
):

    if not rows:
        return

    max_columns = max(
        len(row)
        for row in rows
    )

    if max_columns < 2:
        return

    table = document.add_table(
        rows=len(rows),
        cols=max_columns
    )

    table.style = "Table Grid"

    # ВАЖНО:
    # больше не центрируем саму таблицу.

    table.alignment = (
        WD_TABLE_ALIGNMENT.LEFT
    )

    # Автоматическая ширина

    table.autofit = True

    for row_index, row in enumerate(rows):

        for column_index in range(
            max_columns
        ):

            cell = table.cell(
                row_index,
                column_index
            )

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            if column_index < len(row):

                value = row[
                    column_index
                ].strip()

            else:

                value = ""

            cell.text = value

            for paragraph in cell.paragraphs:

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)

                for run in paragraph.runs:

                    run.font.name = "Arial"
                    run.font.size = Pt(10)

                    # Заголовок таблицы
                    # выделяем только если это
                    # первая строка.

                    if row_index == 0:

                        run.bold = True


# =========================================================
# СОХРАНЕНИЕ WORD
# =========================================================

def save_word_document(
    document
):

    try:

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".docx"
        )

        file_path = temp_file.name

        temp_file.close()

        document.save(
            file_path
        )

        print(
            "✅ WORD СОЗДАН:",
            file_path,
            flush=True
        )

        return file_path

    except Exception as error:

        print(
            "SAVE WORD ERROR:",
            error,
            flush=True
        )

        return None


# =========================================================
# СОЗДАНИЕ WORD
# =========================================================

def create_word_document(
    text,
    ocr_data=None
):

    try:

        print(
            "📄 СОЗДАЁМ WORD — ВИЗУАЛЬНЫЙ РЕЖИМ",
            flush=True
        )

        document = Document()

        # =================================================
        # СТРАНИЦА
        # =================================================

        section = document.sections[0]

        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # =================================================
        # ОСНОВНОЙ ШРИФТ
        # =================================================

        normal_style = document.styles["Normal"]

        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        # =================================================
        # ЕСЛИ НЕТ OCR-КООРДИНАТ
        # =================================================

        if not ocr_data:

            for raw_line in text.splitlines():

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.first_line_indent = Cm(0)

                run = paragraph.add_run(
                    raw_line
                )

                run.font.name = "Arial"
                run.font.size = Pt(11)

            return save_word_document(
                document
            )

        # =================================================
        # OCR ДАННЫЕ
        # =================================================

        lines = ocr_data.get(
            "lines",
            []
        )

        words = ocr_data.get(
            "words",
            []
        )

        if not lines:

            return save_word_document(
                document
            )

        # =================================================
        # РАЗМЕР ИСХОДНОГО ИЗОБРАЖЕНИЯ
        # =================================================

        page_width = max(
            word["right"]
            for word in words
        )

        page_height = max(
            word["bottom"]
            for word in words
        )

        print(
            "📐 OCR PAGE:",
            page_width,
            "x",
            page_height,
            flush=True
        )

        # =================================================
        # РАСЧЁТ МАСШТАБА
        # =================================================

        # Word A4 рабочая ширина примерно 18 см.
        #
        # Все координаты фотографии переводим
        # в координаты Word.

        usable_width_cm = 18.0

        scale_x = (
            usable_width_cm
            /
            max(page_width, 1)
        )

        # =================================================
        # ВЕРТИКАЛЬНАЯ ШКАЛА
        # =================================================

        previous_bottom = 0

        # =================================================
        # ОБРАБОТКА КАЖДОЙ OCR-СТРОКИ
        # =================================================

        for index, line in enumerate(
            lines
        ):

            if not line:
                continue

            text_line = line_to_text(
                line
            )

            if not text_line:
                continue

            # -------------------------------------------------
            # ГРАНИЦЫ СТРОКИ
            # -------------------------------------------------

            left = min(
                word["left"]
                for word in line
            )

            top = min(
                word["top"]
                for word in line
            )

            right = max(
                word["right"]
                for word in line
            )

            bottom = max(
                word["bottom"]
                for word in line
            )

            # -------------------------------------------------
            # ВЫСОТА БУКВ
            # -------------------------------------------------

            heights = [
                word["height"]
                for word in line
                if word["height"] > 0
            ]

            if heights:

                median_height = statistics.median(
                    heights
                )

            else:

                median_height = 20

            # -------------------------------------------------
            # РАЗМЕР ШРИФТА
            # -------------------------------------------------

            font_size = (
                median_height
                * 0.72
            )

            font_size = max(
                7,
                min(
                    18,
                    font_size
                )
            )

            # -------------------------------------------------
            # СОЗДАЁМ ПАРАГРАФ
            # -------------------------------------------------

            paragraph = document.add_paragraph()

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )

            paragraph.paragraph_format.line_spacing = 1.0

            paragraph.paragraph_format.space_after = Pt(0)

            paragraph.paragraph_format.first_line_indent = Cm(0)

            # -------------------------------------------------
            # ГОРИЗОНТАЛЬНАЯ ПОЗИЦИЯ
            # -------------------------------------------------

            left_indent = (
                left
                * scale_x
            )

            # Не даём отступу выйти
            # за пределы страницы.

            left_indent = max(
                0,
                min(
                    16,
                    left_indent
                )
            )

            paragraph.paragraph_format.left_indent = Cm(
                left_indent
            )

            # -------------------------------------------------
            # ВЕРТИКАЛЬНЫЙ ПРОБЕЛ
            # -------------------------------------------------

            if previous_bottom > 0:

                vertical_gap = (
                    top
                    -
                    previous_bottom
                )

                # Преобразуем пиксели
                # в приблизительные points.

                space_before = (
                    vertical_gap
                    * 0.35
                )

                space_before = max(
                    0,
                    min(
                        30,
                        space_before
                    )
                )

                paragraph.paragraph_format.space_before = Pt(
                    space_before
                )

            # -------------------------------------------------
            # ОПРЕДЕЛЯЕМ ЗАГОЛОВОК
            # -------------------------------------------------

            heading = is_heading(
                text_line
            )

            if heading:

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                paragraph.paragraph_format.left_indent = Cm(0)

                font_size = max(
                    13,
                    font_size
                )

                font_size = min(
                    20,
                    font_size
                )

            # -------------------------------------------------
            # СПИСОК
            # -------------------------------------------------

            elif is_numbered_item(
                text_line
            ):

                paragraph.paragraph_format.left_indent = Cm(
                    min(
                        16,
                        left_indent + 0.3
                    )
                )

            elif is_list_item(
                text_line
            ):

                paragraph.paragraph_format.left_indent = Cm(
                    min(
                        16,
                        left_indent + 0.5
                    )
                )

            # -------------------------------------------------
            # ТЕКСТ
            # -------------------------------------------------

            run = paragraph.add_run(
                text_line
            )

            run.font.name = "Arial"

            run.font.size = Pt(
                font_size
            )

            run.bold = heading

            # -------------------------------------------------
            # ЗАПОМИНАЕМ НИЗ СТРОКИ
            # -------------------------------------------------

            previous_bottom = bottom

        # =================================================
        # СОХРАНЕНИЕ
        # =================================================

        print(
            "✅ ВИЗУАЛЬНЫЙ WORD ГОТОВ",
            flush=True
        )

        return save_word_document(
            document
        )

    except Exception as error:

        print(
            "❌ WORD VISUAL ERROR:",
            error,
            flush=True
        )

        return None


# =========================================================
# ОТПРАВКА WORD
# =========================================================

def send_document(
    chat_id,
    file_path,
    filename="document.docx"
):

    try:

        url = f"{API}/sendDocument"

        with open(
            file_path,
            "rb"
        ) as document_file:

            response = requests.post(
                url,
                data={
                    "chat_id": chat_id
                },
                files={
                    "document": (
                        filename,
                        document_file,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                },
                timeout=60
            )

        print(
            "SEND DOCUMENT:",
            response.status_code,
            response.text,
            flush=True
        )

        return response.json()

    except Exception as error:

        print(
            "SEND DOCUMENT ERROR:",
            error,
            flush=True
        )

        return {
            "ok": False
        }


# =========================================================
# КЛАВИАТУРЫ
# =========================================================

def main_menu_keyboard():

    return {
        "inline_keyboard": [

            [
                {
                    "text": "📸 Распознать документ",
                    "callback_data": "ocr"
                }
            ],

            [
                {
                    "text": "📎 Работа с файлом",
                    "callback_data": "file_tools"
                }
            ]

        ]
    }


def ocr_keyboard():

    return {
        "inline_keyboard": [

            [
                {
                    "text": "📄 Создать Word",
                    "callback_data": "create_word"
                }
            ],

            [
                {
                    "text": "🔎 Найти в файле",
                    "callback_data": "find_in_file"
                }
            ],

            [
                {
                    "text": "🌐 Перевести",
                    "callback_data": "translate"
                },

                {
                    "text": "✍️ Улучшить",
                    "callback_data": "improve"
                }
            ],

            [
                {
                    "text": "📸 Распознать ещё",
                    "callback_data": "ocr"
                }
            ],

            [
                {
                    "text": "🏠 Главное меню",
                    "callback_data": "menu"
                }
            ]

        ]
    }


# =========================================================
# START
# =========================================================

def start_command(
    chat_id
):

    send_message(
        chat_id,

        "👋 Привет! Я QEVRA 🚀\n\n"

        "Я могу превратить фотографию документа "
        "в редактируемый Word-файл.\n\n"

        "📸 Распознать документ\n"
        "📄 Восстановить Word\n"
        "🌐 Перевести текст\n"
        "✍️ Улучшить текст\n\n"

        "Просто отправь фотографию документа.",

        main_menu_keyboard()
    )


# =========================================================
# МЕНЮ
# =========================================================

def show_menu(
    chat_id
):

    send_message(
        chat_id,

        "🤖 Что будем делать?",

        main_menu_keyboard()
    )


# =========================================================
# ОБРАБОТКА ФОТО
# =========================================================

def process_photo(
    chat_id,
    user_id,
    message
):

    print(
        "📸 PHOTO PROCESS START",
        flush=True
    )

    now = time.time()

    last_time = photo_processing.get(
        user_id,
        0
    )

    if now - last_time < 10:

        print(
            "⚠️ Фото недавно уже обрабатывалось",
            flush=True
        )

        return

    photo_processing[
        user_id
    ] = now

    send_message(
        chat_id,

        "📸 Фото получил.\n\n"
        "🔎 Анализирую документ..."
    )

    try:

        photos = message.get(
            "photo"
        )

        if not photos:

            send_message(
                chat_id,
                "❌ Фото не найдено."
            )

            return

        # Берём максимальное доступное качество

        photo = photos[-1]

        file_id = photo[
            "file_id"
        ]

        image_bytes = download_telegram_photo(
            file_id
        )

        if image_bytes is None:

            send_message(
                chat_id,

                "❌ Не удалось скачать фотографию."
            )

            return

        # -------------------------------------------------
        # Открываем и готовим изображение
        # -------------------------------------------------

        image = Image.open(
            BytesIO(image_bytes)
        )

        prepared_image = prepare_image(
            image
        )

        # -------------------------------------------------
        # OCR
        # -------------------------------------------------

        ocr_data = recognize_document(
            prepared_image
        )

        if not ocr_data:

            send_message(
                chat_id,

                "😔 Я не смог распознать текст.\n\n"
                "Попробуй сфотографировать документ "
                "при хорошем освещении и без наклона."
            )

            return

        # -------------------------------------------------
        # Текст
        # -------------------------------------------------

        text = ocr_to_text(
            ocr_data
        )

        text = clean_ocr_text(
            text
        )

        if not text:

            send_message(
                chat_id,

                "😔 Текст на фотографии не найден."
            )

            return

        # -------------------------------------------------
        # Сохраняем и текст,
        # и координаты OCR
        # -------------------------------------------------

        user_texts[
            user_id
        ] = text

        user_ocr_data[
            user_id
        ] = ocr_data

        print(
            "✅ OCR СОХРАНЁН",
            "Символов:",
            len(text),
            "Слов:",
            len(
                ocr_data.get(
                    "words",
                    []
                )
            ),
            flush=True
        )

        send_long_message(
            chat_id,

            "📝 Распознанный текст:\n\n"
            + text
        )

        send_message(
            chat_id,

            "🤖 Документ проанализирован.\n\n"
            "Можно восстановить его структуру в Word.",

            ocr_keyboard()
        )

    except Exception as error:

        print(
            "PHOTO ERROR:",
            error,
            flush=True
        )

        send_message(
            chat_id,

            "❌ Ошибка при обработке фотографии."
        )

    finally:

        print(
            "📸 PHOTO PROCESS END",
            flush=True
        )


# =========================================================
# WORD
# =========================================================

def process_create_word(
    chat_id,
    user_id
):

    if user_id not in user_texts:

        send_message(
            chat_id,

            "❌ Сначала отправь фотографию документа."
        )

        return

    text = user_texts[
        user_id
    ]

    ocr_data = user_ocr_data.get(
        user_id
    )

    send_message(
        chat_id,

        "📄 Восстанавливаю документ...\n\n"
        "🔎 Анализирую расположение текста\n"
        "📊 Проверяю таблицы\n"
        "📝 Формирую Word"
    )

    file_path = create_word_document(
        text,
        ocr_data
    )

    if not file_path:

        send_message(
            chat_id,

            "❌ Не удалось создать Word-файл."
        )

        return

    try:

        result = send_document(
            chat_id,
            file_path,
            "QEVRA_document.docx"
        )

        if result.get("ok"):

            send_message(
                chat_id,

                "✅ Готово!\n\n"
                "📄 Документ восстановлен в Word.\n\n"
                "Если что-то получилось неправильно — "
                "пришли фотографию, и мы будем улучшать "
                "распознавание."
            )

        else:

            send_message(
                chat_id,

                "❌ Не удалось отправить Word-файл."
            )

    finally:

        try:

            os.remove(
                file_path
            )

        except Exception:
            pass

# =========================================================
# ПОИСК В ФАЙЛЕ — ЭКСПЕРИМЕНТАЛЬНАЯ ФУНКЦИЯ
# =========================================================

def search_in_file(
    text,
    query
):

    if not text or not query:
        return None

    query = query.strip()

    if not query:
        return None

    # -----------------------------------------------------
    # Поиск без учёта регистра
    # -----------------------------------------------------

    lines = text.splitlines()

    query_lower = query.lower()

    matches = []

    for index, line in enumerate(lines):

        if query_lower in line.lower():

            # Берём несколько соседних строк,
            # чтобы показать контекст.

            start = max(
                0,
                index - 1
            )

            end = min(
                len(lines),
                index + 2
            )

            context = "\n".join(
                lines[start:end]
            ).strip()

            matches.append(
                context
            )

    if not matches:
        return None

    # -----------------------------------------------------
    # Убираем одинаковые результаты
    # -----------------------------------------------------

    unique_matches = []

    for match in matches:

        if match not in unique_matches:

            unique_matches.append(
                match
            )

    # -----------------------------------------------------
    # Ограничиваем количество результатов
    # -----------------------------------------------------

    unique_matches = unique_matches[:10]

    return unique_matches

# =========================================================
# QEVRA DOCUMENT INTELLIGENCE
# УНИВЕРСАЛЬНОЕ ОПРЕДЕЛЕНИЕ ДОКУМЕНТА
# =========================================================

DOCUMENT_DATABASE = {

    "Финансовые документы": {

        "Invoice / Счёт": [
            "invoice",
            "invoice number",
            "amount due",
            "subtotal",
            "total",
            "tax",
            "счёт",
            "сумма",
            "итого"
        ],

        "Receipt / Чек": [
            "receipt",
            "cashier",
            "change",
            "payment",
            "чек",
            "кассир",
            "оплата"
        ],

        "Bank Statement / Банковская выписка": [
            "bank statement",
            "account statement",
            "transaction",
            "balance",
            "debit",
            "credit",
            "банковская выписка",
            "операция",
            "баланс"
        ],

        "Payment Order / Платёжное поручение": [
            "payment order",
            "payment instruction",
            "beneficiary",
            "payer",
            "платёжное поручение",
            "получатель",
            "плательщик"
        ]
    },


    "Юридические документы": {

        "Contract / Договор": [
            "contract",
            "agreement",
            "party",
            "parties",
            "terms and conditions",
            "договор",
            "соглашение",
            "стороны",
            "условия"
        ],

        "Power of Attorney / Доверенность": [
            "power of attorney",
            "authorized representative",
            "attorney",
            "доверенность",
            "представитель",
            "уполномоченный"
        ],

        "Application / Заявление": [
            "application",
            "applicant",
            "request",
            "заявление",
            "заявитель",
            "прошу"
        ]
    },


    "Государственные документы": {

        "Passport / Паспорт": [
            "passport",
            "surname",
            "given name",
            "date of birth",
            "nationality",
            "паспорт",
            "фамилия",
            "имя",
            "дата рождения",
            "гражданство"
        ],

        "Birth Certificate / Свидетельство о рождении": [
            "birth certificate",
            "date of birth",
            "place of birth",
            "свидетельство о рождении",
            "место рождения"
        ],

        "Tax Document / Налоговый документ": [
            "tax",
            "taxpayer",
            "tax authority",
            "налог",
            "налогоплательщик",
            "налоговый"
        ]
    },


    "Образовательные документы": {

        "Diploma / Диплом": [
            "diploma",
            "degree",
            "graduation",
            "university",
            "диплом",
            "образование",
            "университет"
        ],

        "Certificate of Education / Аттестат": [
            "certificate of education",
            "secondary education",
            "school",
            "аттестат",
            "среднее образование",
            "школа"
        ],

        "Transcript / Выписка с оценками": [
            "transcript",
            "grade",
            "course",
            "credits",
            "оценка",
            "дисциплина",
            "кредиты"
        ]
    },


    "Трудовые документы": {

        "Employment Contract / Трудовой договор": [
            "employment contract",
            "employee",
            "employer",
            "salary",
            "position",
            "трудовой договор",
            "работник",
            "работодатель",
            "зарплата",
            "должность"
        ],

        "Job Application / Заявление о приёме": [
            "job application",
            "position",
            "applicant",
            "vacancy",
            "заявление о приёме",
            "должность",
            "вакансия"
        ]
    },


    "Медицинские документы": {

        "Medical Report / Медицинское заключение": [
            "medical report",
            "diagnosis",
            "patient",
            "doctor",
            "medical",
            "медицинское заключение",
            "диагноз",
            "пациент",
            "врач"
        ],

        "Prescription / Рецепт": [
            "prescription",
            "medicine",
            "dosage",
            "doctor",
            "рецепт",
            "лекарство",
            "дозировка",
            "врач"
        ],

        "Laboratory Report / Лабораторный анализ": [
            "laboratory",
            "laboratory result",
            "test result",
            "reference range",
            "результат анализа",
            "лаборатория",
            "референсный диапазон"
        ]
    },


    "Страховые документы": {

        "Insurance Policy / Страховой полис": [
            "insurance policy",
            "policy number",
            "insured",
            "insurance company",
            "страховой полис",
            "номер полиса",
            "застрахованный",
            "страховая компания"
        ],

        "Insurance Claim / Страховой случай": [
            "insurance claim",
            "claim number",
            "incident",
            "страховой случай",
            "номер заявления",
            "происшествие"
        ]
    },


    "Транспортные документы": {

        "Driving Licence / Водительское удостоверение": [
            "driving licence",
            "driver license",
            "driving license",
            "date of birth",
            "expiry",
            "водительское удостоверение",
            "права",
            "срок действия"
        ],

        "Vehicle Registration / Регистрация автомобиля": [
            "vehicle registration",
            "registration number",
            "vehicle",
            "автомобиль",
            "регистрационный номер",
            "регистрация"
        ]
    },


    "Логистические документы": {

        "Bill of Lading / Коносамент": [
            "bill of lading",
            "shipper",
            "consignee",
            "port of loading",
            "port of discharge",
            "cargo",
            "коносамент",
            "грузоотправитель",
            "грузополучатель",
            "груз"
        ],

        "Packing List / Упаковочный лист": [
            "packing list",
            "package",
            "quantity",
            "weight",
            "упаковочный лист",
            "количество",
            "вес"
        ],

        "Delivery Note / Накладная": [
            "delivery note",
            "delivery",
            "goods",
            "quantity",
            "накладная",
            "доставка",
            "товар",
            "количество"
        ]
    },


    "Морские документы": {

        "Muster List / Расписание по тревогам": [
            "muster list",
            "alarm signals",
            "general alarm",
            "abandon ship",
            "fire and emergency",
            "man overboard",
            "flooding",
            "сигналы тревог",
            "оставление судна",
            "человек за бортом",
            "затопление"
        ],

        "Ship Certificate / Судовой сертификат": [
            "ship certificate",
            "certificate of registry",
            "flag state",
            "vessel",
            "судовой сертификат",
            "судно",
            "флаг"
        ]
    },


    "Бизнес документы": {

        "Business Report / Бизнес-отчёт": [
            "business report",
            "financial report",
            "revenue",
            "profit",
            "business",
            "отчёт",
            "выручка",
            "прибыль",
            "бизнес"
        ],

        "Purchase Order / Заказ на покупку": [
            "purchase order",
            "po number",
            "supplier",
            "buyer",
            "заказ на покупку",
            "поставщик",
            "покупатель"
        ]
    },


    "Технические документы": {

        "Technical Manual / Техническое руководство": [
            "manual",
            "installation",
            "operation",
            "maintenance",
            "technical",
            "руководство",
            "установка",
            "эксплуатация",
            "обслуживание",
            "технический"
        ],

        "Specification / Спецификация": [
            "specification",
            "technical specification",
            "model",
            "serial number",
            "спецификация",
            "технические характеристики",
            "модель",
            "серийный номер"
        ]
    },


    "Коммерческие документы": {

        "Quotation / Коммерческое предложение": [
            "quotation",
            "quote",
            "offer",
            "price",
            "commercial offer",
            "коммерческое предложение",
            "цена",
            "предложение"
        ],

        "Order / Заказ": [
            "order",
            "order number",
            "customer",
            "quantity",
            "заказ",
            "номер заказа",
            "клиент",
            "количество"
        ]
    }
}


# =========================================================
# ПОИСК ПРИЗНАКОВ ДОКУМЕНТА
# =========================================================

def identify_document(text):

    if not text:

        return {
            "category": "Неизвестно",
            "type": "Документ не определён",
            "description":
                "В документе не удалось обнаружить текст.",
            "confidence": 0,
            "matched_keywords": []
        }

    text_lower = text.lower()

    best_category = None
    best_type = None
    best_score = 0
    best_keywords = []

    # -----------------------------------------------------
    # Проверяем всю базу
    # -----------------------------------------------------

    for category, documents in DOCUMENT_DATABASE.items():

        for document_type, keywords in documents.items():

            matched = []

            for keyword in keywords:

                if keyword.lower() in text_lower:

                    matched.append(keyword)

            score = len(matched)

            if score > best_score:

                best_score = score
                best_category = category
                best_type = document_type
                best_keywords = matched

    # -----------------------------------------------------
    # Не нашли достаточно признаков
    # -----------------------------------------------------

    if best_score == 0:

        return {
            "category": "Неизвестно",
            "type": "Не удалось точно определить",
            "description":
                "QEVRA не обнаружила достаточного количества "
                "характерных признаков документа.",
            "confidence": 0,
            "matched_keywords": []
        }

    # -----------------------------------------------------
    # Расчёт уверенности
    # -----------------------------------------------------

    if best_score >= 5:

        confidence = 95

    elif best_score == 4:

        confidence = 88

    elif best_score == 3:

        confidence = 75

    elif best_score == 2:

        confidence = 55

    else:

        confidence = 30

    # -----------------------------------------------------
    # Описание
    # -----------------------------------------------------

    descriptions = {

        "Финансовые документы":
            "Документ связан с оплатой, финансами, расчётами или денежными операциями.",

        "Юридические документы":
            "Документ имеет юридическое назначение и может определять права, обязанности или официальные обращения.",

        "Государственные документы":
            "Официальный документ, связанный с государственными органами или удостоверением личности и статуса.",

        "Образовательные документы":
            "Документ связан с образованием, обучением, квалификацией или результатами обучения.",

        "Трудовые документы":
            "Документ связан с трудовыми отношениями, работником, работодателем или трудоустройством.",

        "Медицинские документы":
            "Документ содержит медицинскую информацию, результаты обследований или сведения о лечении.",

        "Страховые документы":
            "Документ связан со страхованием, полисом или страховым случаем.",

        "Транспортные документы":
            "Документ связан с транспортом, водителем или регистрацией транспортного средства.",

        "Логистические документы":
            "Документ связан с перевозкой, грузом, доставкой или упаковкой.",

        "Морские документы":
            "Документ связан с эксплуатацией судна, экипажем, безопасностью или морскими перевозками.",

        "Бизнес документы":
            "Документ используется в деятельности компании или организации.",

        "Технические документы":
            "Документ содержит технические сведения, инструкции, характеристики или требования.",

        "Коммерческие документы":
            "Документ связан с продажей, покупкой, заказом или коммерческим предложением."
    }

    return {
        "category": best_category,
        "type": best_type,
        "description":
            descriptions.get(
                best_category,
                "QEVRA определила назначение документа."
            ),
        "confidence": confidence,
        "matched_keywords": best_keywords
    }


# =========================================================
# ФОРМИРОВАНИЕ ИНФОРМАЦИИ О ДОКУМЕНТЕ
# =========================================================

def format_document_information(text):

    result = identify_document(
        text
    )

    if result["confidence"] == 0:

        return (
            "🧠 АНАЛИЗ ДОКУМЕНТА\n\n"

            "❓ Точный тип определить не удалось.\n\n"

            "QEVRA пока не нашла достаточно признаков "
            "для уверенной классификации этого документа.\n\n"

            "💡 Попробуйте использовать фотографию "
            "с более чётким текстом."
        )

    message = (
        "🧠 АНАЛИЗ ДОКУМЕНТА\n\n"

        f"📂 Категория:\n"
        f"{result['category']}\n\n"

        f"📄 Тип:\n"
        f"{result['type']}\n\n"

        f"📌 Назначение:\n"
        f"{result['description']}\n\n"

        f"🎯 Уверенность:\n"
        f"{result['confidence']}%"
    )

    if result["matched_keywords"]:

        message += (
            "\n\n🔎 Найденные признаки:\n"
        )

        for keyword in result["matched_keywords"][:8]:

            message += (
                f"• {keyword}\n"
            )

    return message


# =========================================================
# УЛУЧШЕНИЕ ТЕКСТА
# =========================================================

def improve_text(
    text
):

    try:

        lines = []

        for line in text.splitlines():

            # Не меняем символы.
            # Только удаляем лишние пробелы.

            line = re.sub(
                r"[ \t]+",
                " ",
                line
            ).strip()

            if line:

                lines.append(
                    line
                )

        return "\n".join(
            lines
        ).strip()

    except Exception as error:

        print(
            "IMPROVE ERROR:",
            error,
            flush=True
        )

        return text


# =========================================================
# ПЕРЕВОД
# =========================================================

def translate_text(
    text
):

    try:

        print(
            "🌐 TRANSLATE START",
            flush=True
        )

        russian = sum(
            1
            for char in text
            if "а" <= char.lower() <= "я"
        )

        english = sum(
            1
            for char in text
            if "a" <= char.lower() <= "z"
        )

        if russian >= english:

            source = "ru"
            target = "en"

        else:

            source = "en"
            target = "ru"

        url = (
            "https://translate.googleapis.com/"
            "translate_a/single"
        )

        params = {

            "client": "gtx",
            "sl": source,
            "tl": target,
            "dt": "t",
            "q": text

        }

        response = requests.get(
            url,
            params=params,
            timeout=20
        )

        if response.status_code != 200:

            return None

        data = response.json()

        result = ""

        for item in data[0]:

            if item and item[0]:

                result += item[0]

        return result.strip()

    except Exception as error:

        print(
            "TRANSLATE ERROR:",
            error,
            flush=True
        )

        return None


# =========================================================
# WEBHOOK
# =========================================================

@app.route(
    "/webhook",
    methods=["POST"]
)
def webhook():

    print(
        "!!! QEVRA WEBHOOK STARTED !!!",
        flush=True
    )

    update = request.get_json(
        silent=True
    )

    if not update:

        return "OK"

    update_id = update.get(
        "update_id"
    )

    if update_id is not None:

        if already_processed(
            update_id
        ):

            return "OK"

    print(
        "UPDATE:",
        update,
        flush=True
    )

    # =====================================================
    # MESSAGE
    # =====================================================

    if "message" in update:

        message = update[
            "message"
        ]

        chat_id = message[
            "chat"
        ][
            "id"
        ]

        user_id = message[
            "from"
        ][
            "id"
        ]

        text = message.get(
            "text",
            ""
        )

        print(
            "MESSAGE:",
            text,
            flush=True
        )

        # -------------------------------------------------
        # START
        # -------------------------------------------------

        if text == "/start":

            start_command(
                chat_id
            )

            return "OK"

        # -------------------------------------------------
        # PHOTO
        # -------------------------------------------------

        if "photo" in message:

            process_photo(
                chat_id,
                user_id,
                message
            )

            return "OK"
            
        # -------------------------------------------------
        # ПОИСК В ФАЙЛЕ
        # -------------------------------------------------

        if text and hasattr(
            webhook,
            "search_mode"
        ) and user_id in webhook.search_mode:

            webhook.search_mode.discard(
                user_id
            )

            document_text = user_texts.get(
                user_id
            )

            if not document_text:

                send_message(
                    chat_id,

                    "❌ Файл не найден."
                )

                return "OK"

            send_message(
                chat_id,

                "🔎 Ищу в документе..."
            )

            results = search_in_file(
                document_text,
                text
            )

            if not results:

                send_message(
                    chat_id,

                    "🔎 Ничего не найдено.\n\n"
                    "Попробуй другое слово или "
                    "более короткую фразу."
                )

            else:

                response = (
                    "🔎 Найдено в документе:\n\n"
                )

                for number, result in enumerate(
                    results,
                    1
                ):

                    response += (
                        f"📌 Результат {number}:\n"
                        f"{result}\n\n"
                    )

                send_long_message(
                    chat_id,
                    response
                )

                send_message(
                    chat_id,

                    "Что дальше?",

                    ocr_keyboard()
                )

            return "OK"
            
    # =====================================================
    # CALLBACK
    # =====================================================

    if "callback_query" in update:

        callback = update[
            "callback_query"
        ]

        callback_id = callback[
            "id"
        ]

        user_id = callback[
            "from"
        ][
            "id"
        ]

        chat_id = callback[
            "message"
        ][
            "chat"
        ][
            "id"
        ]

        callback_data = callback[
            "data"
        ]

        print(
            "CALLBACK:",
            callback_data,
            flush=True
        )

        # -------------------------------------------------
        # OCR
        # -------------------------------------------------

        if callback_data == "ocr":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id
                }
            )

            send_message(
                chat_id,

                "📸 Отправь фотографию документа."
            )
                 # -------------------------------------------------
        # FILE TOOLS
        # -------------------------------------------------

        elif callback_data == "file_tools":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id
                }
            )

            if user_id not in user_texts:

                send_message(
                    chat_id,

                    "📎 Работа с файлом\n\n"
                    "Сначала отправь фотографию документа, "
                    "чтобы QEVRA смогла его распознать."
                )

            else:

                send_message(
                    chat_id,

                    "📎 Работа с файлом\n\n"
                    "Что хочешь сделать?\n\n"
                    "🔎 Найти информацию в документе\n\n"
                    "Экспериментальная функция."
                )

                send_message(
                    chat_id,

                    "Нажми «🔎 Найти в файле» "
                    "или просто выбери нужное действие.",
                    ocr_keyboard()
                )

        # -------------------------------------------------
        # FIND IN FILE
        # -------------------------------------------------

        elif callback_data == "find_in_file":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id
                }
            )

            if user_id not in user_texts:

                send_message(
                    chat_id,

                    "❌ У меня пока нет обработанного файла.\n\n"
                    "Сначала отправь фотографию документа."
                )

            else:

                send_message(
                    chat_id,

                    "🔎 Поиск в файле включён.\n\n"
                    "Напиши, что нужно найти.\n\n"
                    "Например:\n"
                    "• номер судна\n"
                    "• капитан\n"
                    "• дату\n"
                    "• телефон\n"
                    "• сумму\n"
                    "• слово или фразу\n\n"
                    "Экспериментальная функция."
                )

                # Флаг ожидания поискового запроса

                if not hasattr(
                    webhook,
                    "search_mode"
                ):

                    webhook.search_mode = set()

                webhook.search_mode.add(
                    user_id
                )
                
        # -------------------------------------------------
        # CREATE WORD
        # -------------------------------------------------

        elif callback_data == "create_word":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id,
                    "text": "📄 Восстанавливаю документ..."
                }
            )

            process_create_word(
                chat_id,
                user_id
            )

        # -------------------------------------------------
        # TRANSLATE
        # -------------------------------------------------

        elif callback_data == "translate":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id,
                    "text": "🌐 Перевожу..."
                }
            )

            if user_id not in user_texts:

                send_message(
                    chat_id,

                    "❌ Нет текста для перевода.\n\n"
                    "Сначала отправь фотографию."
                )

            else:

                text = user_texts[
                    user_id
                ]

                send_message(
                    chat_id,
                    "🌐 Перевожу..."
                )

                translated = translate_text(
                    text
                )

                if translated:

                    send_long_message(
                        chat_id,

                        "🌐 Перевод:\n\n"
                        + translated
                    )

                    send_message(
                        chat_id,

                        "Готово. Что дальше?",

                        ocr_keyboard()
                    )

                else:

                    send_message(
                        chat_id,

                        "❌ Не удалось выполнить перевод."
                    )

        # -------------------------------------------------
        # IMPROVE
        # -------------------------------------------------

        elif callback_data == "improve":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id,
                    "text": "✍️ Улучшаю..."
                }
            )

            if user_id not in user_texts:

                send_message(
                    chat_id,

                    "❌ Сначала отправь фотографию."
                )

            else:

                text = user_texts[
                    user_id
                ]

                improved = improve_text(
                    text
                )

                user_texts[
                    user_id
                ] = improved

                send_long_message(
                    chat_id,

                    "✍️ Улучшенный текст:\n\n"
                    + improved
                )

                send_message(
                    chat_id,

                    "Готово. Что дальше?",

                    ocr_keyboard()
                )

        # -------------------------------------------------
        # MENU
        # -------------------------------------------------

        elif callback_data == "menu":

            telegram(
                "answerCallbackQuery",
                {
                    "callback_query_id": callback_id
                }
            )

            show_menu(
                chat_id
            )

    return "OK"


# =========================================================
# HOME
# =========================================================

@app.route(
    "/",
    methods=["GET"]
)
def home():

    return "QEVRA is alive! 🚀"


# =========================================================
# START SERVER
# =========================================================

if __name__ == "__main__":

    port = int(
        os.environ.get(
            "PORT",
            10000
        )
    )

    print(
        "🚀 QEVRA запускается...",
        flush=True
    )

    app.run(
        host="0.0.0.0",
        port=port
    )
